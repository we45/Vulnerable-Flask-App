from flask import session, Flask, jsonify, request, Response, render_template, render_template_string, url_for
from flask_sqlalchemy import SQLAlchemy
import jwt
from jwt.exceptions import DecodeError, MissingRequiredClaimError, InvalidKeyError
import json
import hashlib
import datetime
import os
from faker import Faker
import random
from werkzeug.utils import secure_filename
from docx import Document
import yaml

from tornado.wsgi import WSGIContainer
from tornado.httpserver import HTTPServer
from tornado.ioloop import IOLoop
import os
import base64

app_port = os.environ.get('APP_PORT', 5050)


app = Flask(__name__, template_folder='templates')
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///test.db'
app.config['SECRET_KEY_HMAC'] = 'secret'
app.config['SECRET_KEY_HMAC_2'] = 'am0r3C0mpl3xK3y'
app.secret_key = 'F12Zr47j\3yX R~X@H!jmM]Lwf/,?KT'
app.config['STATIC_FOLDER'] = None

db = SQLAlchemy(app)

class User(db.Model):
    id = db.Column(db.Integer, primary_key = True)
    username = db.Column(db.String(80), unique = True)
    password = db.Column(db.String(80), unique = True)


    def __repr__(self):
        return "<User {0}>".format(self.username)

class Customer(db.Model):
    id = db.Column(db.Integer, primary_key = True)
    first_name = db.Column(db.String(80))
    last_name = db.Column(db.String(80))
    email = db.Column(db.String(80))
    ccn = db.Column(db.String(80), nullable = True)
    username = db.Column(db.String(80))
    password = db.Column(db.String(150))


    def __repr__(self):
        return "<User {0} {1}>".format(self.first_name, self.last_name)

@app.before_first_request
def setup_users():

    db.create_all()

    if not User.query.first():
        user = User()
        user.username = 'admin'
        user.password = 'admin123'
        db.session.add(user)
        db.session.commit()
    if not Customer.query.first():
        for i in range(0,5):
            fake = Faker()
            cust = Customer()
            cust.first_name = fake.first_name()
            cust.last_name = fake.last_name()
            cust.email = fake.simple_profile(sex = None)['mail']
            cust.username = fake.simple_profile(sex = None)['username']
            cust.password = str(base64.b64encode(os.urandom(16)))
            cust.ccn = fake.credit_card_number(card_type=None)
            db.session.add(cust)
            db.session.commit()


def get_exp_date():
    exp_date = datetime.datetime.utcnow() + datetime.timedelta(minutes = 240)
    return exp_date

def verify_jwt(token):
    try:
        decoded = jwt.decode(token, app.config['SECRET_KEY_HMAC'], verify=True, issuer = 'we45', leeway=10, algorithms=['HS256'])
        print("JWT Token from API: {0}".format(decoded))
        return True
    except DecodeError:
        print("Error in decoding token")
        return False
    except MissingRequiredClaimError as e:
        print('Claim required is missing: {0}'.format(e))
        return False

def insecure_verify(token):
    decoded = jwt.decode(token, verify = False)
    print(decoded)
    return True

@app.errorhandler(404)
def pnf(e):
    template = '''<html>
    <head>
    <title>Error</title>
    </head>
    <body>
    <h1>Oops that page doesn't exist!!</h1>
    <h3>%s</h3>
    </body>
    </html>
    ''' % request.url

    return render_template_string(template, dir = dir, help = help, locals = locals),404

def has_no_empty_params(rule):
    default = rule.defaults if rule.defaults is not None else ()
    arguments = rule.arguments if rule.arguments is not None else ()
    return len(default) >= len(arguments)

@app.route('/', methods = ['GET'])
def sitemap():
    links = []
    for rule in app.url_map.iter_rules():
        print(rule)
        if ("GET" in rule.methods or "POST" in rule.methods) and has_no_empty_params(rule):
            if not 'static' in rule.endpoint:
                url = url_for(rule.endpoint, **(rule.defaults or {}))
                links.append((url, rule.endpoint, ','.join(rule.methods)))

    return render_template('index.html', urls = links)


@app.route('/register/user', methods = ['POST'])
def reg_customer():
    try:
        content = request.json
        if content:
            username = content['username']
            password = content['password']
            hash_pass = hashlib.md5(password).hexdigest()
            new_user = User(username, hash_pass)
            db.session.add(new_user)
            db.session.commit()
            user_created = 'User: {0} has been created'.format(username)
            return jsonify({'Created': user_created}),200
    except Exception as e:
        return jsonify({'Error': str(e.message)}),404

@app.route('/register/customer', methods = ['POST'])
def reg_user():
    try:
        content = request.json
        if content:
            username = content['username']
            password = content['password']
            first_name = content['first_name']
            last_name = content['last_name']
            email = content['email']
            ccn = content['ccn']
            cust = Customer(first_name, last_name, email, username, password, ccn)
            db.session.add(cust)
            db.session.commit()
            user_created = 'Customer: {0} has been created'.format(username)
            return jsonify({'Created': user_created}),200
    except Exception as e:
        return jsonify({'Error': str(e.message)}),404


@app.route('/login', methods = ['POST'])
def login():
    '''
    You will need to authenticate to this URI first. You will need to pass a JSON body with a username and password key.
    If you enter a valid username and password, a JWT token is returned in the HTTP Response in the Authorization header.
    This token can be used for subsequent requests.
    '''
    try:
        content = request.json
        print(content)
        username = content['username']
        password = content['password']
        auth_user = User.query.filter_by(username = username, password = password).first()
        if auth_user:
            auth_token = jwt.encode({'user': username, 'exp': get_exp_date(), 'nbf': datetime.datetime.utcnow(), 'iss': 'we45', 'iat': datetime.datetime.utcnow()}, app.config['SECRET_KEY_HMAC'], algorithm='HS256')
            resp = Response(json.dumps({'Authenticated': True, "User": username}))
            #resp.set_cookie('SESSIONID', auth_token)
            resp.headers['Authorization'] = "{0}".format(auth_token)
            resp.status_code = 200
            resp.mimetype = 'application/json'
            return resp
        else:
            return jsonify({'Error': 'No User here...'}),404
    except:
        return jsonify({'Error': 'Unable to recognize Input'}),404

@app.route('/fetch/customer', methods = ['POST'])
def fetch_customer():
    token = request.headers.get('Authorization')
    if not token:
        return jsonify({'Error': 'Not Authenticated!'}),403
    else:
        if not verify_jwt(token):
            return jsonify({'Error': 'Invalid Token'}),403
        else:
            content = request.json
            if content:
                customer_id = content['id']
                customer_record = Customer.query.get(customer_id)
                customer_dict = {'id': customer_record.id, 'firstname': customer_record.first_name,
                                 'lastname': customer_record.last_name, 'email': customer_record.email,
                                 'cc_num': customer_record.ccn, 'username': customer_record.username
                                }
                if customer_record:
                    return jsonify(customer_dict),200
                else:
                    return jsonify({'Error': 'No Customer Found'}),404
            else:
                return jsonify({'Error': 'Invalid Request'}),400


@app.route('/get/<cust_id>', methods = ['GET'])
def get_customer(cust_id):
    token = request.headers.get('Authorization')
    if not token:
        return jsonify({'Error': 'Not Authenticated!'}), 403
    else:
        if not insecure_verify(token):
            return jsonify({'Error': 'Invalid Token'}), 403
        else:
            if cust_id:
                customer_record = Customer.query.get(cust_id)
                if customer_record:
                    customer_dict = {'id': customer_record.id, 'firstname': customer_record.first_name,
                                 'lastname': customer_record.last_name, 'email': customer_record.email,
                                 'cc_num': customer_record.ccn, 'username': customer_record.username
                                }
                    return jsonify(customer_dict),200
                else:
                    return jsonify({'Error': 'No Customer Found'}),404
            else:
                return jsonify({'Error': 'Invalid Request'}),400




@app.route('/search', methods = ['POST'])
def search_customer():
    token = request.headers.get('Authorization')
    if not token:
        return jsonify({'Error': 'Not Authenticated!'}),403
    else:
        if not verify_jwt(token):
            return jsonify({'Error': 'Invalid Token'}),403
        else:
            content = request.json
            results = []
            if content:
                try:
                    search_term = content['search']
                    print(search_term)
                    str_query = "SELECT first_name, last_name, username FROM customer WHERE username = '%s';" % search_term
                    # mycust = Customer.query.filter_by(username = search_term).first()
                    # return jsonify({'Customer': mycust.username, 'First Name': mycust.first_name}),200

                    search_query = db.engine.execute(str_query)
                    for result in search_query:
                        results.append(list(result))
                    print(results)
                    return jsonify(results),200
                except Exception as e:
                    template = '''<html>
                        <head>
                        <title>Error</title>
                        </head>
                        <body>
                        <h1>Oops Error Occurred</h1>
                        <h3>%s</h3>
                        </body>
                        </html>
                        ''' % str(e)
                    return render_template_string(template, dir=dir, help=help, locals=locals), 404


@app.route("/xxe")
def index():
    return render_template(
        'test.html')


@app.route("/xxe_uploader", methods=['GET', 'POST'])  # /<string:name>/")
def hello():
    if request.method == 'POST':

        f = request.files['file']
        rand = random.randint(1, 100)
        fname = secure_filename(f.filename)
        fname = str(rand) + fname  # change file name
        cwd = os.getcwd()
        file_path = cwd + '/Files/' + fname
        f.save(file_path)  # save file locally

        # Access saved file
        document = Document(file_path)
        for para in document.paragraphs:
            print (para.text)  # '\n\n'.join([para.text for paragraph in document.paragraphs])

    # return "file uploaded successfully"
    return render_template('view.html', name=para.text)

@app.route("/yaml")
def yaml_upload():
    return render_template(
        'yaml_test.html')

@app.route("/yaml_hammer", methods = ['POST'])
def yaml_hammer():
    if request.method == "POST":
        f = request.files['file']
        rand = random.randint(1, 100)
        fname = secure_filename(f.filename)
        fname = str(rand) + fname  # change file name
        cwd = os.getcwd()
        file_path = cwd + '/Files/' + fname
        f.save(file_path)  # save file locally

        with open(file_path, 'r') as yfile:
            y = yfile.read()

        ydata = yaml.load(y)

    return render_template('view.html', name = json.dumps(ydata))



if __name__ == "__main__":
    http_server = HTTPServer(WSGIContainer(app))
    http_server.listen(app_port)
    IOLoop.instance().start()
    # app.run(debug = True, host = '0.0.0.0', port = app_port)
